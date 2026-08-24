"""
trait_research.py — Moteur de recherche de gènes candidats par thème,
généralisé à n'importe quelle espèce/thème du dataset (pas seulement
verse/quinoa). Conçu pour être appelé directement depuis app.py via un
nouvel onglet Streamlit "Recherche par thème".

Pipeline :
  1. Le sujet tapé par l'utilisateur ("verse chez le quinoa", "sécheresse
     chez le maïs") est matché contre une bibliothèque de modèles
     pré-écrits (TOPIC_TEMPLATES). Si aucun modèle ne correspond, un
     message clair l'indique (extension future : génération de mots-clés
     via un LLM pour les sujets non couverts).
  2. search_candidates() scanne le fichier <espece>_all_sources.json pour
     les mots-clés du modèle, restreint aux gènes origin=sequence_backed.
  3. score_candidates() trie par spécificité + diversité de catégories.
  4. fetch_pubmed_references() interroge l'API NCBI E-utilities (gratuite,
     sans clé) pour chaque gène retenu.
  5. generate_docx_report() produit un .docx téléchargeable directement
     depuis l'app.

Usage direct (hors Streamlit, pour test) :
    python trait_research.py --species-file path.json --topic "verse"
"""

from __future__ import annotations

import json
import re
import time
import argparse
import unicodedata
from pathlib import Path
from dataclasses import dataclass, field
from urllib.parse import quote

import requests


def _normalize_text(s: str) -> str:
    """Normalise unicode (NFC) et casse avant toute comparaison de string
    tapée par un utilisateur.

    Sans ça : un accent tapé via un terminal, un heredoc PowerShell, ou
    collé depuis certaines sources peut être encodé en forme décomposée
    (NFD -- 'e' + accent combinant séparé) plutôt que composée (NFC --
    un seul point de code pour 'é'). Visuellement identique à l'écran,
    mais `"secheresse" in q` échoue silencieusement entre les deux
    formes, sans la moindre erreur. C'est ce qui a fait échouer
    match_topic("sécheresse") testé en ligne de commande."""
    return unicodedata.normalize("NFC", s or "").strip().lower()

# ── Bibliothèque de modèles de thèmes ───────────────────────────────────────
# Chaque modèle définit des catégories de mots-clés (utilisées pour le score
# de diversité) et une répartition tier A (signal fort) / tier B (signal
# faible) pour le score de spécificité. NOISY = mots-clés volontairement
# exclus car trop génériques (bruit constaté empiriquement sur le quinoa).

SPECIES_ALIASES: dict[str, str] = {
    "quinoa": "chenopodium quinoa",
    "riz": "oryza sativa", "rice": "oryza sativa",
    "mais": "zea mays", "maïs": "zea mays", "maize": "zea mays", "corn": "zea mays",
    "tomate": "solanum lycopersicum", "tomato": "solanum lycopersicum",
    "raisin": "vitis vinifera", "grape": "vitis vinifera", "vigne": "vitis vinifera",
    "tabac": "nicotiana tabacum", "tobacco": "nicotiana tabacum",
    "pomme de terre": "solanum tuberosum", "potato": "solanum tuberosum",
}


def resolve_species_filter(species_input: str | None) -> str | None:
    """Normalise un nom d'espèce tapé par l'utilisateur (commun ou
    scientifique, n'importe quelle langue courante) vers le nom
    scientifique réellement stocké dans le champ "organism" du dataset.
    Sans ça, un nom commun ("maize") ne matche jamais "Zea mays" et
    filtre silencieusement TOUS les gènes -- c'est exactement le bug
    rencontré en testant le thème sécheresse sur le maïs (0 candidats)."""
    if not species_input:
        return None
    key = _normalize_text(species_input)
    return SPECIES_ALIASES.get(key, key)  # si pas dans la table, on tente tel quel


TOPIC_TEMPLATES: dict[str, dict] = {
    "verse": {
        "aliases": ["verse", "lodging", "tige", "rigidite"],
        "label": "Verse / résistance de la tige (lodging)",
        "pubmed_context": "lodging",
        "keywords": {
            "lignification": [
                "lignin", "lignification", "cell wall", "secondary cell wall",
                "cellulose synthase", "laccase", "peroxidase",
                "cinnamyl alcohol", "phenylpropanoid", "4cl", "ccoaomt", "comt",
            ],
            "rigidite_tige": [
                "stem", "culm", "stalk", "lodging", "verse", "internode",
                "mechanical strength", "stem strength",
            ],
            "hormonal": [
                "gibberellin", "gibberellic", "della", "ga20ox", "ga3ox",
                "ga2ox", "brassinosteroid", "bri1",
            ],
        },
        "tier_a": {
            "lignin", "lignification", "laccase", "cellulose synthase",
            "secondary cell wall", "phenylpropanoid", "cinnamyl alcohol",
            "4cl", "comt", "ccoaomt", "della", "gibberellin", "gibberellic",
            "ga20ox", "ga3ox", "ga2ox", "brassinosteroid", "bri1",
            "verse", "lodging", "stalk", "mechanical strength",
            "stem strength",
        },
        "tier_b": {"peroxidase", "culm", "internode"},
        "noisy": {"stem", "cell wall"},
    },
    "secheresse": {
        "aliases": ["secheresse", "sécheresse", "drought", "hydrique", "eau"],
        "label": "Tolérance à la sécheresse / stress hydrique",
        "pubmed_context": "drought",
        "keywords": {
            "signalisation_aba": [
                "abscisic acid", "aba receptor", "pyr/pyl", "snrk2",
                "aba signaling",
            ],
            "reponse_osmotique": [
                "osmotic stress", "proline", "dehydrin", "late embryogenesis",
                "lea protein", "osmotic adjustment",
            ],
            "regulation_stomates": [
                "stomatal closure", "guard cell", "aquaporin", "water use efficiency",
            ],
            "facteurs_transcription": [
                "dreb", "nac transcription factor", "wrky", "myb drought",
            ],
        },
        "tier_a": {
            "abscisic acid", "aba receptor", "pyr/pyl", "snrk2",
            "dehydrin", "lea protein", "dreb", "aquaporin",
            "stomatal closure",
        },
        "tier_b": {
            "aba signaling", "osmotic stress", "proline",
            "late embryogenesis", "osmotic adjustment", "guard cell",
            "water use efficiency", "nac transcription factor", "wrky",
            "myb drought",
        },
        "noisy": set(),
    },
}


def match_topic(user_query: str) -> str | None:
    """Trouve le modèle de thème le plus proche du texte tapé par
    l'utilisateur. Retourne la clé du modèle, ou None si aucun match --
    dans ce cas l'app doit indiquer que le thème n'est pas encore couvert
    (extension future : génération de mots-clés via LLM)."""
    q = _normalize_text(user_query)
    for key, tpl in TOPIC_TEMPLATES.items():
        if any(_normalize_text(alias) in q for alias in tpl["aliases"]):
            return key
    return None


# ── Recherche de candidats (généralisée depuis search_lodging_candidates_v2.py) ──

def _collect_text(gene: dict) -> str:
    parts = []

    def add(v):
        if isinstance(v, str):
            parts.append(v.lower())
        elif isinstance(v, dict):
            for x in v.values():
                add(x)
        elif isinstance(v, list):
            for x in v:
                add(x)

    for field_name in ("description", "common_name", "traits"):
        add(gene.get(field_name))
    annotation = gene.get("annotation") or {}
    for field_name in ("go_terms", "kegg_pathways", "mapman", "tf_family"):
        add(annotation.get(field_name))
    add(gene.get("literature"))
    return " | ".join(parts)


def _has_real_sequence(gene: dict) -> bool:
    seq = gene.get("sequence") or {}
    return bool(seq.get("dna") or seq.get("rna") or seq.get("protein"))


def load_genes(path: Path) -> dict:
    with open(path, "r", encoding="utf-8") as f:
        raw = json.load(f)
    if isinstance(raw, dict) and "genes" in raw:
        raw = raw["genes"]
    if isinstance(raw, dict) and all(isinstance(v, dict) for v in raw.values()):
        return raw
    if isinstance(raw, list):
        return {g.get("gene_id", str(i)): g for i, g in enumerate(raw)}
    raise ValueError("Format de fichier non reconnu")


def search_candidates(genes: dict, template: dict, species_filter: str | None = None) -> list[dict]:
    all_kw = [(kw, cat) for cat, kws in template["keywords"].items() for kw in kws]
    noisy_tagged = {f"{cat}:{kw}" for cat, kws in template["keywords"].items()
                     for kw in kws if kw in template["noisy"]}

    results = []
    for gene_id, gene in genes.items():
        if not isinstance(gene, dict):
            continue
        if gene.get("origin") != "sequence_backed":
            continue
        if species_filter:
            organism = str(gene.get("organism", "")).lower()
            if species_filter.lower() not in organism:
                continue
        text = _collect_text(gene)
        matches = {f"{cat}:{kw.strip()}" for kw, cat in all_kw if kw in text} - noisy_tagged
        if not matches or not _has_real_sequence(gene):
            continue
        results.append({"gene_id": gene_id, "gene": gene, "matches": matches})
    return results


def score_candidates(candidates: list[dict], template: dict) -> list[dict]:
    tier_a, tier_b = template["tier_a"], template["tier_b"]
    for c in candidates:
        kws = {m.split(":", 1)[1] for m in c["matches"]}
        n_tier_a = len(kws & tier_a)
        n_tier_b = len(kws & tier_b)
        n_categories = len({m.split(":")[0] for m in c["matches"]})
        score = n_tier_a * 3 + n_tier_b * 1 + max(0, n_categories - 1) * 2
        c["score"] = score
        c["n_categories"] = n_categories
        c["categories"] = sorted({m.split(":")[0] for m in c["matches"]})
    candidates.sort(key=lambda c: c["score"], reverse=True)
    return candidates


# ── Sourcing PubMed automatique (NCBI E-utilities, gratuit, sans clé) ──────

PUBMED_BASE = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils"


def fetch_pubmed_references(gene_name: str, extra_terms: list[str] | None = None,
                             retmax: int = 3, sleep: float = 0.34) -> list[dict]:
    """Cherche jusqu'à `retmax` publications PubMed pour un gène + contexte.
    `sleep` respecte la limite de ~3 req/s de l'API NCBI sans clé."""
    terms = [gene_name] + (extra_terms or [])
    query = " AND ".join(f'"{t}"[Title/Abstract]' for t in terms)

    try:
        search_resp = requests.get(
            f"{PUBMED_BASE}/esearch.fcgi",
            params={"db": "pubmed", "term": query, "retmax": retmax, "retmode": "json"},
            timeout=15,
        )
        pmids = search_resp.json().get("esearchresult", {}).get("idlist", [])
        if not pmids:
            return []
        time.sleep(sleep)

        summary_resp = requests.get(
            f"{PUBMED_BASE}/esummary.fcgi",
            params={"db": "pubmed", "id": ",".join(pmids), "retmode": "json"},
            timeout=15,
        )
        result = summary_resp.json().get("result", {})
        refs = []
        for pmid in pmids:
            doc = result.get(pmid, {})
            if not doc:
                continue
            refs.append({
                "pmid": pmid,
                "title": doc.get("title", ""),
                "journal": doc.get("fulljournalname", ""),
                "year": (doc.get("pubdate", "") or "")[:4],
                "url": f"https://pubmed.ncbi.nlm.nih.gov/{pmid}/",
            })
        time.sleep(sleep)
        return refs
    except Exception:
        return []  # une source qui échoue ne doit jamais casser tout le run


# ── Export Word (python-docx, à intégrer directement dans l'app) ──────────

def generate_docx_report(topic_label: str, species: str, candidates: list[dict],
                          out_path: str, top_n: int = 30) -> str:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_heading(f"Gènes candidats — {topic_label} chez {species}", level=1)
    p = doc.add_paragraph()
    p.add_run(
        f"Rapport généré automatiquement par Plant Gene Analyzer. "
        f"{len(candidates)} candidats trouvés, top {min(top_n, len(candidates))} affichés."
    ).italic = True

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Gène", "Accession", "Score", "Catégories", "Référence(s)"]):
        hdr[i].text = h

    for c in candidates[:top_n]:
        row = table.add_row().cells
        row[0].text = c["gene"].get("common_name", "") or c["gene_id"]
        row[1].text = c["gene_id"]
        row[2].text = str(c["score"])
        row[3].text = ", ".join(c["categories"])
        refs = c.get("references", [])
        row[4].text = "\n".join(f"{r['title']} ({r['year']}) — {r['url']}" for r in refs) or "Non trouvée"

    doc.save(out_path)
    return out_path


# ── Point d'entrée CLI pour tester le module hors Streamlit ────────────────



def _looks_like_systematic_id(name: str) -> bool:
    """Détecte un identifiant systématique (ex. 'Zm00001e014008', une
    accession UniProt brute) plutôt qu'un vrai nom de gène lisible --
    ces identifiants sont inutiles comme terme de recherche PubMed
    (ils n'apparaissent presque jamais tels quels dans un titre/résumé)."""
    if not name:
        return True
    if re.match(r"^[A-Za-z]{1,3}\d{4,}[A-Za-z0-9]*$", name):  # Zm00001e014008
        return True
    if re.match(r"^[A-Z][0-9][A-Z0-9]{3,8}$", name):  # accession UniProt brute (P41979...)
        return True
    return False


def pick_search_term(gene: dict, gene_id: str) -> str | None:
    """Choisit le meilleur terme de recherche PubMed disponible pour un
    gène : nom de gène lisible en priorité, sinon un terme GO
    (molecular_function/biological_process en priorité, plus
    informatifs que cellular_component), sinon la dernière portion
    d'une description mapman, sinon abandon."""
    common_name = (gene.get("common_name") or "").split(":")[0].strip()
    if common_name and not _looks_like_systematic_id(common_name):
        return common_name

    annotation = gene.get("annotation") or {}
    go_terms = annotation.get("go_terms") or []
    for pref_aspect in ("molecular_function", "biological_process"):
        for gt in go_terms:
            if isinstance(gt, dict) and gt.get("aspect") == pref_aspect and gt.get("term"):
                return gt["term"]
    if go_terms and isinstance(go_terms[0], dict) and go_terms[0].get("term"):
        return go_terms[0]["term"]

    mapman = annotation.get("mapman") or []
    if mapman and isinstance(mapman[0], dict) and mapman[0].get("description"):
        # ex. "Cell wall organisation.lignin.monolignol conjugation..." -> dernier segment
        return mapman[0]["description"].split(".")[-1].strip()

    return None


# ── Interface Streamlit (à appeler depuis app.py, voir guide d'intégration) ──

try:
    import streamlit as st
except ImportError:
    st = None

SPECIES_FILES: dict[str, str] = {
    "quinoa": "chenopodium_quinoa_all_sources.json",
    "riz": "oryza_sativa_all_sources.json",
    "maïs": "zea_mays_all_sources.json",
    "tomate": "solanum_lycopersicum_all_sources.json",
    "raisin": "vitis_vinifera_all_sources.json",
    "tabac": "nicotiana_tabacum_all_sources.json",
    "pomme de terre": "solanum_tuberosum_all_sources.json",
}


def render_trait_research_tab(species_dir: str) -> None:
    """Section Streamlit complète "Recherche par thème". À appeler depuis
    app.py avec le chemin du dossier contenant les fichiers
    <espece>_all_sources.json (voir guide d'intégration).

    Ne dépend d'AUCUNE séquence saisie par l'utilisateur -- fonctionne de
    façon totalement indépendante du flux d'analyse de séquence existant.
    """
    if st is None:
        raise RuntimeError("streamlit n'est pas installé dans cet environnement.")

    st.markdown("### 🌱 Recherche de gènes candidats par thème")
    st.markdown(
        "Choisis une espèce et décris un problème agronomique "
        "(ex. *verse chez le quinoa*, *sécheresse chez le maïs*) pour obtenir "
        "une liste de gènes candidats sourcée dans la littérature scientifique."
    )

    col1, col2 = st.columns([1, 2])
    with col1:
        species_label = st.selectbox("Espèce", options=list(SPECIES_FILES.keys()))
    with col2:
        topic_query = st.text_input(
            "Thème / problème étudié",
            placeholder="ex. verse, sécheresse, résistance au froid...",
        )

    fetch_refs = st.checkbox(
        "Chercher les références PubMed (plus lent, ~0.7s par gène)",
        value=True,
    )
    top_n = st.slider("Nombre de candidats à afficher", 5, 50, 20)

    if not st.button("🔍 Lancer la recherche", type="primary"):
        return

    topic_key = match_topic(topic_query) if topic_query else None
    if not topic_key:
        st.warning(
            f"⚠ Thème non reconnu. Thèmes actuellement disponibles : "
            f"{', '.join(t['label'] for t in TOPIC_TEMPLATES.values())}. "
            f"Pour ajouter un nouveau thème, voir TOPIC_TEMPLATES dans trait_research.py."
        )
        return

    template = TOPIC_TEMPLATES[topic_key]
    species_file = Path(species_dir) / SPECIES_FILES[species_label]

    with st.spinner(f"Chargement des données {species_label}..."):
        genes = _load_genes_cached(str(species_file))

    species_filter = resolve_species_filter(species_label)
    candidates = search_candidates(genes, template, species_filter)
    candidates = score_candidates(candidates, template)

    if not candidates:
        st.info("Aucun candidat trouvé pour cette combinaison espèce/thème.")
        return

    st.success(f"{len(candidates)} candidats trouvés — top {min(top_n, len(candidates))} affichés.")

    if fetch_refs:
        progress = st.progress(0, text="Recherche PubMed en cours...")
        for i, c in enumerate(candidates[:top_n]):
            term = pick_search_term(c["gene"], c["gene_id"])
            c["references"] = fetch_pubmed_references(term, extra_terms=[template["pubmed_context"]]) if term else []
            progress.progress((i + 1) / min(top_n, len(candidates)))
        progress.empty()
    else:
        for c in candidates[:top_n]:
            c["references"] = []

    table_rows = [{
        "Gène": c["gene"].get("common_name", "") or c["gene_id"],
        "Accession": c["gene_id"],
        "Score": c["score"],
        "Catégories": ", ".join(c["categories"]),
        "Références": len(c.get("references", [])),
    } for c in candidates[:top_n]]
    st.dataframe(table_rows, width="stretch")

    with st.expander("Voir le détail des références PubMed trouvées"):
        for c in candidates[:top_n]:
            refs = c.get("references", [])
            if refs:
                name = c["gene"].get("common_name", "") or c["gene_id"]
                st.markdown(f"**{name}** ({c['gene_id']})")
                for r in refs:
                    st.markdown(f"- {r['title']} ({r['year']}) — [{r['pmid']}]({r['url']})")

    out_path = f"/tmp/rapport_{topic_key}_{species_label}.docx"
    generate_docx_report(template["label"], species_label, candidates, out_path, top_n)
    with open(out_path, "rb") as f:
        st.download_button(
            "📄 Télécharger le rapport Word",
            data=f.read(),
            file_name=f"candidats_{topic_key}_{species_label}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


def _load_genes_cached(species_file: str) -> dict:
    """Wrapper cache Streamlit autour de load_genes -- évite de relire
    et re-parser un fichier de plusieurs dizaines de milliers de gènes à
    chaque interaction avec un widget de la page."""
    if st is not None:
        cached = st.cache_data(show_spinner=False)(load_genes)
        return cached(Path(species_file))
    return load_genes(Path(species_file))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--species-file", required=True)
    ap.add_argument("--topic", required=True)
    ap.add_argument("--species-name", default=None)
    ap.add_argument("--fetch-refs", action="store_true", help="Interroge PubMed (lent, ~0.35s/gène)")
    ap.add_argument("--top-n", type=int, default=15)
    ap.add_argument("--out", default="rapport_candidats.docx")
    args = ap.parse_args()

    topic_key = match_topic(args.topic)
    if not topic_key:
        print(f"⚠ Thème '{args.topic}' non couvert par la bibliothèque de modèles.")
        print(f"  Thèmes disponibles : {list(TOPIC_TEMPLATES.keys())}")
        return

    template = TOPIC_TEMPLATES[topic_key]
    genes = load_genes(Path(args.species_file))
    species_filter = resolve_species_filter(args.species_name)
    candidates = search_candidates(genes, template, species_filter)
    candidates = score_candidates(candidates, template)

    print(f"{len(candidates)} candidats trouvés pour '{template['label']}'.\n")
    for c in candidates[:args.top_n]:
        print(f"  [{c['score']:2d}] {c['gene_id']} — {c['gene'].get('common_name', '')} ({', '.join(c['categories'])})")

    if args.fetch_refs:
        print("\nRecherche PubMed en cours...")
        for c in candidates[:args.top_n]:
            term = pick_search_term(c["gene"], c["gene_id"])
            if not term:
                c["references"] = []
                print(f"  {c['gene_id']}: pas de terme de recherche exploitable, ignoré")
                continue
            c["references"] = fetch_pubmed_references(term, extra_terms=[template["pubmed_context"]])
            print(f"  {term}: {len(c['references'])} référence(s)")

    out = generate_docx_report(template["label"], args.species_name or "?", candidates, args.out, args.top_n)
    print(f"\n✓ Rapport écrit : {out}")


if __name__ == "__main__":
    main()